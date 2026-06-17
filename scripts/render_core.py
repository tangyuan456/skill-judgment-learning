#!/usr/bin/env python3
"""
render_core.py — 自包含三格式渲染器（md/docx/html），不依赖上游 skill。

提供：
  - render_md(report_data) -> str
  - render_docx(report_data, output_path, charts_abs_dir)
  - render_html(report_data) -> str

report_data 结构：
{
    "cover": {"title": "...", "subtitle": "...", "date": "...", "version": "..."},
    "sections": [
        {
            "id": "ch1", "title": "第1章 ...",
            "blocks": [...],           # 顶层 block
            "subsections": [           # 子节
                {
                    "id": "ch1_1", "title": "1.1 ...",
                    "blocks": [
                        {"type": "paragraph", "text": "..."},
                        {"type": "table", "headers": [...], "rows": [[...], ...]},
                        {"type": "image", "path": "charts/xxx.png", "caption": "..."},
                        {"type": "insight", "level": "L1/L2/L3", "text": "...", "source": "..."},
                        {"type": "bullet", "items": ["...", "..."]},
                    ]
                }
            ]
        }
    ]
}
"""
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ==============================
# Markdown 渲染
# ==============================

def render_md(report_data: Dict[str, Any]) -> str:
    lines = []
    cover = report_data.get('cover', {})

    # 封面
    lines.append(f"# {cover.get('title', '数据库性能测试报告')}")
    lines.append('')
    if cover.get('subtitle'):
        lines.append(f"> {cover['subtitle']}")
        lines.append('')
    lines.append(f"- **日期**：{cover.get('date', datetime.now().strftime('%Y-%m-%d'))}")
    lines.append(f"- **版本**：{cover.get('version', 'v1.0')}")
    lines.append('')
    lines.append('---')
    lines.append('')

    # 目录
    lines.append('## 目录')
    lines.append('')
    for sec in report_data.get('sections', []):
        lines.append(f"- [{sec['title']}](#{_md_anchor(sec['title'])})")
        for sub in sec.get('subsections', []):
            lines.append(f"  - [{sub['title']}](#{_md_anchor(sub['title'])})")
    lines.append('')
    lines.append('---')
    lines.append('')

    # 各章节
    for sec in report_data.get('sections', []):
        lines.append(f"## {sec['title']}")
        lines.append('')

        # 顶层 blocks
        _render_md_blocks(sec.get('blocks', []), lines)

        # 子节
        for sub in sec.get('subsections', []):
            lines.append(f"### {sub['title']}")
            lines.append('')
            _render_md_blocks(sub.get('blocks', []), lines)

    return '\n'.join(lines)


def _render_md_blocks(blocks: List[Dict], lines: List[str]):
    for b in blocks:
        typ = b.get('type', 'paragraph')
        if typ == 'paragraph':
            lines.append(b.get('text', ''))
            lines.append('')
        elif typ == 'table':
            headers = b.get('headers', [])
            rows = b.get('rows', [])
            # 表头
            lines.append('| ' + ' | '.join(str(h) for h in headers) + ' |')
            lines.append('| ' + ' | '.join('---' for _ in headers) + ' |')
            for row in rows:
                lines.append('| ' + ' | '.join(str(c) for c in row) + ' |')
            lines.append('')
        elif typ == 'image':
            caption = b.get('caption', '')
            path = b.get('path', '')
            lines.append(f'![{caption}]({path})')
            if caption:
                lines.append(f'*{caption}*')
            lines.append('')
        elif typ == 'insight':
            level = b.get('level', '')
            text = b.get('text', '')
            lines.append(f'> **[{level}]** {text}')
            lines.append('')
        elif typ == 'bullet':
            for item in b.get('items', []):
                lines.append(f'- {item}')
            lines.append('')


def _md_anchor(title: str) -> str:
    return re.sub(r'[^a-zA-Z0-9\u4e00-\u9fff-]', '', title.lower().replace(' ', '-'))


# ==============================
# HTML 渲染
# ==============================

_HTML_CSS = """
<style>
  body { font-family: 'Microsoft YaHei', 'PingFang SC', 'Hiragino Sans GB', Arial, sans-serif; max-width: 960px; margin: 0 auto; padding: 20px; color: #333; }
  h1 { text-align: center; font-size: 2em; margin-bottom: 4px; }
  h1 + p { text-align: center; color: #666; font-size: 1.1em; }
  h2 { border-bottom: 2px solid #1f77b4; padding-bottom: 6px; margin-top: 32px; }
  h3 { margin-top: 24px; color: #1f77b4; }
  table { border-collapse: collapse; width: 100%; margin: 12px 0; }
  th { background: #D9E2F3; padding: 8px 12px; text-align: left; border: 1px solid #bbb; }
  td { padding: 6px 12px; border: 1px solid #ddd; }
  tr:nth-child(even) td { background: #f9f9f9; }
  .cover { text-align: center; padding: 40px 0; border-bottom: 1px solid #ddd; margin-bottom: 24px; }
  .cover .meta { color: #888; margin-top: 8px; }
  .toc { background: #f5f5f5; padding: 16px 24px; border-radius: 6px; margin: 16px 0; }
  .toc ul { list-style: none; padding-left: 0; }
  .toc li { margin: 4px 0; }
  .toc a { color: #1f77b4; text-decoration: none; }
  img { max-width: 100%; height: auto; display: block; margin: 12px auto; }
  .caption { text-align: center; color: #666; font-size: 0.9em; margin-bottom: 16px; }
  .insight { padding: 10px 16px; margin: 8px 0; border-left: 4px solid #1f77b4; background: #f0f5ff; }
  .insight.L1 { border-left-color: #2ca02c; background: #f0fff0; }
  .insight.L2 { border-left-color: #ff7f0e; background: #fff8f0; }
  .insight.L3 { border-left-color: #9467bd; background: #f8f0ff; }
  .insight .level { font-weight: bold; margin-right: 8px; }
  ul.bullet { padding-left: 20px; }
  ul.bullet li { margin: 4px 0; }
</style>
"""


def render_html(report_data: Dict[str, Any]) -> str:
    cover = report_data.get('cover', {})
    sections = report_data.get('sections', [])

    # 生成 TOC
    toc_items = []
    body_items = []

    for sec in sections:
        sec_id = sec.get('id', '')
        toc_items.append(f'<li><a href="#{sec_id}">{sec["title"]}</a></li>')
        body = [f'<h2 id="{sec_id}">{sec["title"]}</h2>']
        body += _render_html_blocks(sec.get('blocks', []))

        for sub in sec.get('subsections', []):
            sub_id = sub.get('id', '')
            body.append(f'<h3 id="{sub_id}">{sub["title"]}</h3>')
            body += _render_html_blocks(sub.get('blocks', []))

        body_items.append('\n'.join(body))

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{cover.get('title', '数据库性能测试报告')}</title>
{_HTML_CSS}
</head>
<body>

<div class="cover">
  <h1>{cover.get('title', '数据库性能测试报告')}</h1>
  {f'<p>{cover["subtitle"]}</p>' if cover.get('subtitle') else ''}
  <p class="meta">日期：{cover.get('date', '')} &nbsp;|&nbsp; 版本：{cover.get('version', 'v1.0')}</p>
</div>

<div class="toc">
  <h3>目录</h3>
  <ul>
    {''.join(toc_items)}
  </ul>
</div>

{''.join(f'<div class="section">{s}</div>' for s in body_items)}

</body>
</html>"""
    return html


def _render_html_blocks(blocks: List[Dict]) -> List[str]:
    result = []
    for b in blocks:
        typ = b.get('type', 'paragraph')
        if typ == 'paragraph':
            result.append(f'<p>{b.get("text", "")}</p>')
        elif typ == 'table':
            headers = b.get('headers', [])
            rows = b.get('rows', [])
            html = '<table><thead><tr>'
            html += ''.join(f'<th>{h}</th>' for h in headers)
            html += '</tr></thead><tbody>'
            for row in rows:
                html += '<tr>' + ''.join(f'<td>{c}</td>' for c in row) + '</tr>'
            html += '</tbody></table>'
            result.append(html)
        elif typ == 'image':
            path = b.get('path', '')
            caption = b.get('caption', '')
            result.append(f'<img src="{path}" alt="{caption}">')
            if caption:
                result.append(f'<p class="caption">{caption}</p>')
        elif typ == 'insight':
            level = b.get('level', '')
            text = b.get('text', '')
            result.append(f'<div class="insight {level}"><span class="level">[{level}]</span>{text}</div>')
        elif typ == 'bullet':
            items = b.get('items', [])
            result.append('<ul class="bullet">' + ''.join(f'<li>{item}</li>' for item in items) + '</ul>')
    return result


# ==============================
# DOCX 渲染
# ==============================

def render_docx(report_data: Dict[str, Any], output_path: str, charts_abs_dir: Optional[str] = None):
    """渲染 Word 文档。"""
    if not HAS_DOCX:
        raise ImportError('需要安装 python-docx: pip install python-docx')

    doc = Document()
    _docx_setup_styles(doc)
    cover = report_data.get('cover', {})

    # 封面
    _docx_add_cover(doc, cover)
    doc.add_page_break()

    # 各章节
    for sec in report_data.get('sections', []):
        doc.add_heading(sec['title'], level=1)

        _docx_render_blocks(doc, sec.get('blocks', []), charts_abs_dir)

        for sub in sec.get('subsections', []):
            doc.add_heading(sub['title'], level=2)
            _docx_render_blocks(doc, sub.get('blocks', []), charts_abs_dir)

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)


def _docx_setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _docx_add_cover(doc, cover):
    for _ in range(4):
        doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = 1  # center
    run = title.add_run(cover.get('title', '数据库性能测试报告'))
    run.font.size = Pt(26)
    run.font.bold = True

    if cover.get('subtitle'):
        sub = doc.add_paragraph()
        sub.alignment = 1
        run = sub.add_run(cover['subtitle'])
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')
    meta = doc.add_paragraph()
    meta.alignment = 1
    meta.add_run(f"日期：{cover.get('date', '')}   版本：{cover.get('version', 'v1.0')}")


def _docx_render_blocks(doc, blocks, charts_abs_dir):
    for b in blocks:
        typ = b.get('type', 'paragraph')
        if typ == 'paragraph':
            doc.add_paragraph(b.get('text', ''))
        elif typ == 'table':
            headers = b.get('headers', [])
            rows = b.get('rows', [])
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 表头
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = str(h)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
            # 数据行
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    table.rows[ri + 1].cells[ci].text = str(val)
            doc.add_paragraph('')
        elif typ == 'image':
            path = b.get('path', '')
            caption = b.get('caption', '')
            if charts_abs_dir and not os.path.isabs(path):
                full = os.path.join(charts_abs_dir, os.path.basename(path))
            elif not os.path.isabs(path):
                full = path
            else:
                full = path
            if os.path.exists(full):
                doc.add_picture(full, width=Inches(5.5))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif typ == 'insight':
            level = b.get('level', '')
            text = b.get('text', '')
            p = doc.add_paragraph()
            run = p.add_run(f'[{level}] {text}')
            run.font.italic = True
            if level == 'L1':
                run.font.color.rgb = RGBColor(0x2c, 0xa0, 0x2c)
            elif level == 'L2':
                run.font.color.rgb = RGBColor(0xff, 0x7f, 0x0e)
        elif typ == 'bullet':
            for item in b.get('items', []):
                doc.add_paragraph(item, style='List Bullet')
